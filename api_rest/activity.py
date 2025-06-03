import sqlite3
import requests
from docx import Document

def collect_data_country(name):
    url = f'https://restcountries.com/v3.1/name/{name}'
    response = requests.get(url)
    if response.status_code != 200:
      print(f'Erro ao consultar o país {name}')
      return None
    
    datas = response.json()
    info = datas[0]
    try:
       coin = list(info.get('currencies',{}).values())[0]
       idioma = list(info.get('languages',{}).values())[0]
    except:
       coin = {'name':'Desconhecida','symbol':'?'}
       idioma = {"Idioma desconhecido"}
    return {
       "nome_comum": info.get('name',{}).get('common','N/A'),
       'nome_oficial':info.get('name',{}).get('official'),
       'capital':info.get('capital',['n/a'])[0],
       'continente':info.get('continents','n/a'),
       "regiao":info.get('region','n/a'),
       'sub_regiao':info.get('subregion','n/a'),
       'população':info.get('population',0),
       'area':info.get('area',0),
       'moeda':coin.get('name','n/a'),
       'simbolo_moeda':coin.get('symbol','?'),
       'idioma':idioma,
       'fuso_horario':info.get('timezones',['N/A'])[0],
       'bandeira':info.get('flags',{}).get('png','')
    }

def save_bd(data):
   conn = sqlite3.connect('dados_paises.db')
   cursor = conn.cursor()

   cursor.execute('''CREATE TABLE paises(
                  nome_oficial TEXT,nome_comum TEXT, capital TEXT,
                  continente TEXT, regiao TEXT, sub_regiao TEXT, população INTEGER, area REAL, moeda TEXT,
                  simbolo_moeda TEXT, idioma TEXT, fuso_horario TEXT,bandeira TEXT
                  )
                   ''')
   
   """  cursor.execute('''
       INSERT INTO paises VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)
''',tuple(data.values()))"""
   
   conn.commit()
   conn.close()
   

def gerar_relatorio(data_list,aluno='Gabriel'):
   doc = Document()
   doc.add_heading('Relatório de dados geográficos',0)
   for data in data_list:
      doc.add_heading(data['nome comum'],level=1)
      for key,value in data.items():
         doc.add_paragraph(f'{key.replace('_',' ').capitalize()}: {value}')
      doc.add_paragraph('')

      file_name = f'relatório_{file_name}.docx'
      doc.save(file_name)
      return file_name

def main():
   data_list = []
   names = ['brazil','japan']
   for i in range(2):
      name = names[i]
      data = collect_data_country(name)
      if data:
         save_bd(data)
         data_list.append(data)
        
   if data_list:
      file_name = gerar_relatorio(data_list)
      print(f'Relatório gerado: {file_name}')

if __name__ == '__main__':
   main()