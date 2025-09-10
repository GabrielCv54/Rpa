from docx import Document

"""#1- criar um novo documento
document = Document()

#2 - Adicionar um título
document.add_heading("relatori",level=1)

#3- Adicionar parágrafo
document.add_paragraph("Este documento descreve o arquivo gerado durante a aula de rpa")

#4- Adicionar subtitulo
document.add_heading('Atividades',level=2)

#5 - Adicionando uma lista de atividades
activity = [
    "Reunião com a equipe ",
    "Desenvolvimento de novos módulos",
    "Testes de Funcionalidades",
    "Treinamento para novos colaboradores"
    ]

for act in activity:
    document.add_paragraph(act,style="List Bullet")

# 6 - Adicionar um subtítulo para considerações finais
document.add_heading('Considerações finais',level=2)

#7 - Parágrafo final com as considerações
document.add_paragraph('Todas as metas estabelecidas foram cumpridas')

#8- Salvar
document.save('relatór.docx')
print('Salvamento feito com sucesso!!')"""


doc = Document('geradores/arquivos_word/relatório.docx')
for p in doc.paragraphs:
    print(p.text)

table = doc.add_table(rows=3,cols=3)
table.cell(0,0).text = "Funcionário"
table.cell(0,1).text = "Profissão"
table.cell(0,2).text = "Salário"
table.cell(1,0).text = 'Maria Guedes'
table.cell(1,1).text = 'Analista de rh'
table.cell(1,2).text = 'R$ 3500,00'
table.cell(2,0).text = 'Renato Lopes'
table.cell(2,1).text = 'Corretor de Imóveis'
table.cell(2,2).text = 'R$ 3541,00'
doc.save('geradores/arquivos_word/relatório_com_tabela.docx')